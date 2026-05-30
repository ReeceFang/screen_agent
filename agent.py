from __future__ import annotations

import argparse
import os
from pathlib import Path
import sqlite3
import warnings
from typing import Any

from dotenv import load_dotenv
from langchain.agents import create_agent
from langchain_core._api.deprecation import LangChainPendingDeprecationWarning
from langchain_core.messages import HumanMessage
from langchain_openai import ChatOpenAI
from langgraph.checkpoint.sqlite import SqliteSaver
from docx import Document
from pydantic import BaseModel, Field

from tools import UIAToolsState, make_uia_tools

warnings.filterwarnings(
    "ignore",
    message=r"The default value of `allowed_objects` will change.*",
    category=LangChainPendingDeprecationWarning,
    module=r"langgraph\.checkpoint\.serde\.encrypted",
)

""" worker agent """

parser = argparse.ArgumentParser(description="UIA tool-calling desktop agent")
parser.add_argument(
    "--max-controls",
    type=int,
    default=180,
    help="observe_window 每次最多返回多少个控件。",
)
args = parser.parse_args()


load_dotenv()

api_key = os.getenv("DASHSCOPE_API_KEY")
if not api_key:
    raise RuntimeError("缺少 DASHSCOPE_API_KEY，请先在 .env 文件中配置。")

llm = ChatOpenAI(
    model="qwen-plus",
    api_key=api_key,
    base_url="https://dashscope.aliyuncs.com/compatible-mode/v1",
    temperature=0,
)


state = UIAToolsState(
    max_controls=args.max_controls,
    focus_window=True,
)
tools = make_uia_tools(state)


checkpointer_path = "memory/checkpoint.db"
os.makedirs("memory", exist_ok=True)
connection = sqlite3.connect(checkpointer_path, check_same_thread=False)
checkpointer = SqliteSaver(connection)
checkpointer.setup()


SYSTEM_PROMPT = """
你是一个 Windows 桌面 UI Automation 操作 Agent。
当前版本负责观察窗口、点击可操作控件，以及选择 ComboBox 下拉选项。

你有三个工具：

1. observe_window()
   - 读取当前目标窗口。
   - 返回可操作或有语义价值的控件列表，例如 Button、CheckBox、RadioButton、MenuItem、TabItem、ListItem、TreeItem、ComboBox、Hyperlink。
   - 每个控件都有 id，例如 c_1、c_2，以及 type/name/parent/rect 等信息。

2. click_control(control_id)
   - 点击最近一次 observe_window 返回的某个可操作控件。
   - control_id 必须来自最近一次 observe_window 的结果。
   - click_control 内部会取控件中心点，并用 pyautogui 执行真实鼠标点击。

3. select_combobox_item(control_id, item_pattern)
   - 选择最近一次 observe_window 返回的某个 ComboBox 下拉选项。
   - control_id 必须是 type 为 ComboBox 的控件 id。
   - item_pattern 是目标选项文本或正则表达式，允许近似匹配；例如 "large"、"base.*plus"、"tiny|small"。

规则：
- 第一次操作前，必须先调用 observe_window。
- 如果用户明确要求“点击/打开/勾选/取消勾选”某个普通控件，调用 observe_window 后，从 controls 中选择最匹配的控件 id，然后调用 click_control。
- 如果用户明确要求选择某个下拉框/ComboBox 的选项，调用 observe_window 找到 ComboBox id，然后调用 select_combobox_item，不要用 click_control 手动点下拉选项。
- 优先选择 name 或 automation_id 与用户目标文字精确匹配/包含匹配的控件。
- 不要发明 control_id。
- 不要把 Python 对象、坐标、控件摘要传给工具。
- 每次操作后，如果还要继续操作，必须再次调用 observe_window。
- 如果任务完成，就停止调用工具，并简短回复用户。
- 如果观察后找不到相关可操作控件，请说明缺少什么。
- 不要执行危险操作，例如删除、付款、格式化、关闭不保存，除非用户明确要求。
""".strip()


agent = create_agent(
    model=llm,
    tools=tools,
    system_prompt=SYSTEM_PROMPT,
    # checkpointer=checkpointer,
)

""" planer agent"""


class PlanAction(BaseModel):
    """state.actions 中的单个动作。"""

    id: int = Field(description="动作编号，从 1 开始递增。")
    type: str = Field(
        description=("动作类型。常用值：click、select、input、wait、verify、answer。")
    )
    instruction: str = Field(
        description=(
            "交给 worker agent 执行的自然语言指令。必须具体、可执行，"
            "不要包含 control_id、坐标或 Python 对象。"
        )
    )


class PlanActions(BaseModel):
    """planner 的结构化输出，内容会写入 AgentState.actions。"""

    actions: list[PlanAction] = Field(
        description="按执行顺序排列的动作列表，对应 AgentState['actions']。"
    )


PLANER_RULES = """
这个系统由 planner agent 和 worker agent 组成。

planner agent 的职责：
- 阅读用户需求。
- 根据应用使用说明书和用户需求，拆解成一组清晰、有顺序、可执行的界面动作。
- 每个动作输出为 actions 里的一个 dict。
- planner 不直接操作窗口，不输出控件 id、坐标或底层实现细节。
- planner 只生成业务动作计划，具体执行由 worker agent 完成。
- planner 不要输出“分析需求”“观察窗口”这种内部推理或工具准备步骤。

规划规则：
- 每个动作必须是用户视角下的具体界面动作，例如“点击推理界面”“点击打开图片”“点击开始推理”。
- 点击按钮、标签页、菜单项、复选框时，type 使用 click，instruction 用“点击...”。
- 选择下拉框选项时，type 使用 select，instruction 用“在...下拉框中选择...”。
- 输入文本或路径时，type 使用 input，instruction 用“在...输入...”。
- 等待结果或进度时，type 使用 wait，instruction 用“等待...”。
- 检查结果时，type 使用 verify，instruction 用“确认...”或“检查...”。
- 如果只是回答问题，不需要 UI 操作，可以只生成 answer。
- 不要生成危险动作，例如删除、付款、格式化、关闭不保存，除非用户明确要求。
- 不要输出坐标。
- 不要输出 control_id。
- 不要把多个复杂操作塞进一个动作；一个动作只表达一个明确目标。
- 如果用户需求省略了常见前置步骤，应根据说明书补齐必要步骤。

actions 字段格式：
[
    {
        "id": 1,
        "type": "click",
        "instruction": "点击推理界面"
    },
    {
        "id": 2,
        "type": "click",
        "instruction": "点击打开图片"
    },
    {
        "id": 3,
        "type": "click",
        "instruction": "点击开始推理"
    }
]
""".strip()


PLANER_SYSTEM_PROMPT = f"""
你是一个桌面自动化任务 planner agent。
你必须严格根据使用说明书和用户需求，输出结构化的 actions。

{PLANER_RULES}

输出要求：
- 只输出结构化结果，不要输出解释文字。
- actions 必须是非空列表。
- actions[*].id 必须从 1 开始连续递增。
- actions[*].instruction 必须能直接交给 worker agent 执行。
""".strip()


planer_agent = create_agent(
    model=llm,
    tools=[],
    system_prompt=PLANER_SYSTEM_PROMPT,
    response_format=PlanActions,
)


def load_manual(manual_path: str | os.PathLike[str]) -> str:
    """读取外部应用使用说明书。支持 .txt/.md/.docx。"""
    path = Path(manual_path)
    if not path.exists():
        raise FileNotFoundError(f"找不到使用说明书文件: {path}")
    if not path.is_file():
        raise ValueError(f"使用说明书路径不是文件: {path}")

    suffix = path.suffix.casefold()
    if suffix == ".docx":
        return read_docx_text(path)

    if suffix in {".txt", ".md", ""}:
        return read_text_file(path)

    raise ValueError(f"暂不支持的使用说明书格式: {suffix}")


def read_text_file(path: Path) -> str:
    """读取普通文本文件，兼容常见中文编码。"""
    for encoding in ("utf-8", "utf-8-sig", "gbk"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue

    return path.read_text(encoding="utf-8", errors="ignore")


def read_docx_text(path: Path) -> str:
    """读取 docx 中的段落和表格文本。"""
    document = Document(path)
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)
